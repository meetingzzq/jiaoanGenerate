"""
Word文档操作模块 - 处理Word文档的读取、填充和保存
"""
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_text(cell, text: str):
    """
    设置单元格文本，保持原有字体格式，并设置左对齐和垂直居中
    彻底清除原有内容，确保无缩进
    """
    # 处理文本：去除首尾空格、删除多余空行、确保顶格
    text = str(text).strip()
    text = re.sub(r'\n{2,}', '\n', text)  # 删除多余空行
    text = re.sub(r'^[ \t]+', '', text, flags=re.MULTILINE)  # 删除行首空格
    
    # 获取第一个段落的字体设置（用于保持格式）
    font_name = None
    font_size = None
    if cell.paragraphs and cell.paragraphs[0].runs:
        font_name = cell.paragraphs[0].runs[0].font.name
        font_size = cell.paragraphs[0].runs[0].font.size
    
    # 删除所有段落
    while len(cell.paragraphs) > 0:
        p_element = cell.paragraphs[-1]._element
        p_element.getparent().remove(p_element)
    
    # 创建全新的段落
    p = cell.add_paragraph()
    
    # 重置段落格式，确保顶格无缩进
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.left_indent = None
    pf.first_line_indent = None
    pf.space_before = None
    pf.space_after = None
    
    # 添加新文本
    run = p.add_run(text)
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = font_size
    
    # 设置单元格垂直居中
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cover_cell_text(cell, text: str):
    """
    设置封面表格单元格文本，文本居中，两边下划线充满单元格
    用于院系、授课班级等封面信息填充
    """
    # 处理文本
    text = str(text).strip()
    
    # 获取原有字体设置（从其他已填充的单元格获取参考格式）
    font_name = None
    font_size = None
    if cell.paragraphs and cell.paragraphs[0].runs:
        font_name = cell.paragraphs[0].runs[0].font.name
        font_size = cell.paragraphs[0].runs[0].font.size
    
    # 删除所有段落，重新创建
    while len(cell.paragraphs) > 0:
        p_element = cell.paragraphs[-1]._element
        p_element.getparent().remove(p_element)
    
    # 创建全新的段落
    p = cell.add_paragraph()
    
    # 设置段落格式 - 居中对齐
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.left_indent = None
    pf.first_line_indent = None
    pf.space_before = None
    pf.space_after = None
    
    # 添加左侧下划线
    left_underline = p.add_run("_" * 15)
    left_underline.font.underline = WD_UNDERLINE.SINGLE
    if font_name:
        left_underline.font.name = font_name
    if font_size:
        left_underline.font.size = font_size
    
    # 添加空格
    p.add_run(" ")
    
    # 添加文本（无下划线）
    text_run = p.add_run(text)
    if font_name:
        text_run.font.name = font_name
    if font_size:
        text_run.font.size = font_size
    
    # 添加空格
    p.add_run(" ")
    
    # 添加右侧下划线
    right_underline = p.add_run("_" * 15)
    right_underline.font.underline = WD_UNDERLINE.SINGLE
    if font_name:
        right_underline.font.name = font_name
    if font_size:
        right_underline.font.size = font_size
    
    # 设置单元格垂直居中
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def find_row_index_by_keyword(table, keyword: str) -> int:
    """
    在表格中查找包含关键词的行索引（支持模糊匹配，忽略空格）
    """
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            cell_text_normalized = cell.text.replace(" ", "").replace("\n", "")
            keyword_normalized = keyword.replace(" ", "")
            if keyword_normalized in cell_text_normalized:
                return i
    return -1


def clear_old_process_rows(process_table):
    """
    清除教学实施过程表格中的旧教学环节行
    保留表头行和课外作业、教学反思行
    """
    while True:
        homework_idx = find_row_index_by_keyword(process_table, "课外作业")
        if homework_idx <= 2:
            break
        row_to_delete = homework_idx - 1
        process_table._element.remove(process_table.rows[row_to_delete]._element)


def insert_process_steps(process_table, process_steps: list):
    """
    在教学实施过程表格中插入新的教学环节
    """
    for step in reversed(process_steps):
        header_row = process_table.rows[1]._element
        new_row = process_table.add_row()
        header_row.addnext(new_row._element)
        
        set_cell_text(new_row.cells[0], f"{step['环节']}（{step['时间']}）")
        set_cell_text(new_row.cells[1], step['内容'])
        set_cell_text(new_row.cells[2], step['教师活动'])
        set_cell_text(new_row.cells[3], step['学生活动'])


class LessonPlanDoc:
    """教案文档类，封装Word文档操作"""
    
    def __init__(self, template_path: str):
        self.doc = Document(template_path)
        if len(self.doc.tables) < 3:
            raise ValueError("模板表格数量不足！需要3个表格：基础信息+教案内容+教学实施过程")
        
        self.info_table = self.doc.tables[0]
        self.content_table = self.doc.tables[1]
        self.process_table = self.doc.tables[2]
    
    def fill_basic_info(self, course_info: dict):
        """填充基础信息表格（封面表格）"""
        # 保留模板原来的样子，不进行填充
        pass
    
    def fill_content_info(self, course_info: dict):
        """填充教案内容表格的基础信息部分"""
        # 第0行: 课题名称
        set_cell_text(self.content_table.cell(0, 1), course_info["课题名称"])
        # 第1行: 授课班级(列1), 授课地点(列5)
        set_cell_text(self.content_table.cell(1, 1), course_info["授课班级"])
        set_cell_text(self.content_table.cell(1, 5), course_info["授课地点"])
        # 第2行: 授课时间(列1), 授课学时(列3), 授课类型(列5)
        set_cell_text(self.content_table.cell(2, 1), course_info["授课时间"])
        set_cell_text(self.content_table.cell(2, 3), course_info["授课学时"])
        set_cell_text(self.content_table.cell(2, 5), course_info["授课类型"])
    
    def fill_content_module(self, row: int, text: str):
        """填充教案内容表格的某个模块"""
        set_cell_text(self.content_table.cell(row, 1), text)
    
    def fill_process_table(self, process_steps: list, homework_text: str):
        """填充教学实施过程表格"""
        # 清除旧的教学环节行
        clear_old_process_rows(self.process_table)
        
        # 插入新的教学环节
        insert_process_steps(self.process_table, process_steps)
        
        # 填充课外作业和教学反思（教学反思保持空白）
        set_cell_text(self.process_table.cell(len(self.process_table.rows)-2, 1), homework_text)
        set_cell_text(self.process_table.cell(len(self.process_table.rows)-1, 1), "")
    
    def save(self, output_path: str):
        """保存文档"""
        self.doc.save(output_path)
